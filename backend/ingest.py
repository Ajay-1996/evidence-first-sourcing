"""Deterministic pre-parsing: artifact file → evidence text the model reads.

The goal is faithful, locator-rich text — sheet+cell for Excel, page markers for PDF,
paragraph/table indices for Word, labelled raw text for email/csv. Images pass through
untouched for vision. No interpretation happens here.
"""
from __future__ import annotations

import base64
from pathlib import Path

FILES_DIR = Path(__file__).resolve().parent.parent / "data" / "vendor_files"


def _unreadable(path: Path, err: Exception) -> str:
    return (f"(THIS DOCUMENT COULD NOT BE PARSED: {type(err).__name__}. "
            f"It may be password-protected, encrypted, or corrupted. "
            f"File '{path.name}' was received but NO content is extractable — "
            f"report it as received-but-unreadable; do not invent any values.)")


def xlsx_to_text(path: Path) -> str:
    from openpyxl import load_workbook
    try:
        wb = load_workbook(path, data_only=False)   # formulas as written
        wbv = load_workbook(path, data_only=True)   # last cached values
    except Exception as e:
        return _unreadable(path, e)
    out = []
    for ws in wb.worksheets:
        state = "" if ws.sheet_state == "visible" else f" [{ws.sheet_state.upper()}]"
        merged = ws.merged_cells.ranges
        out.append(f"=== sheet: {ws.title}{state} ===")
        if merged:
            out.append(f"(merged ranges: {', '.join(str(m) for m in merged)})")
        wsv = wbv[ws.title]
        for row in ws.iter_rows():
            cells = []
            for c in row:
                if c.value is None:
                    continue
                if isinstance(c.value, str) and c.value.startswith("="):
                    cached = wsv[c.coordinate].value
                    cells.append(f"{c.coordinate}={c.value!r}⇒cached:{cached!r}")
                else:
                    cells.append(f"{c.coordinate}={c.value!r}")
            if cells:
                out.append("  ".join(cells))
    return "\n".join(out)


def pdf_to_text(path: Path) -> str:
    import pdfplumber
    out = []
    try:
        with pdfplumber.open(path) as pdf:
            for i, page in enumerate(pdf.pages, 1):
                out.append(f"=== page {i} ===")
                out.append(page.extract_text() or "(no extractable text on this page)")
    except Exception as e:
        return _unreadable(path, e)
    return "\n".join(out)


def docx_to_text(path: Path) -> str:
    import docx
    d = docx.Document(path)
    out = []
    for i, p in enumerate(d.paragraphs, 1):
        if p.text.strip():
            out.append(f"[para {i}] {p.text}")
    for ti, t in enumerate(d.tables, 1):
        out.append(f"=== table {ti} ===")
        for ri, row in enumerate(t.rows, 1):
            out.append(f"[t{ti} row {ri}] " + " | ".join(c.text for c in row.cells))
    return "\n".join(out)


def image_to_block(path: Path) -> dict:
    media = "image/jpeg" if path.suffix.lower() in (".jpg", ".jpeg") else "image/png"
    return {"type": "image",
            "source": {"type": "base64", "media_type": media,
                       "data": base64.standard_b64encode(path.read_bytes()).decode()}}


def artifact_to_content(filename: str):
    """Returns ('text', str) or ('image', block) for one artifact file."""
    path = FILES_DIR / filename
    suf = path.suffix.lower()
    if suf == ".xlsx":
        return "text", f"--- {filename} (Excel) ---\n{xlsx_to_text(path)}"
    if suf == ".pdf":
        return "text", f"--- {filename} (PDF) ---\n{pdf_to_text(path)}"
    if suf == ".docx":
        return "text", f"--- {filename} (Word) ---\n{docx_to_text(path)}"
    if suf in (".jpg", ".jpeg", ".png"):
        return "image", image_to_block(path)
    return "text", f"--- {filename} ---\n{path.read_text()}"
