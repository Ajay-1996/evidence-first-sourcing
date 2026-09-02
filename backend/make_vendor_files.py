"""Fabricate the five supplier artifacts. Five formats, five failure modes:

  ALPHA   alpha_quote.xlsx        clean 30/30 control, questionnaire complete
  BOXCO   boxco_quote.docx        27/30, terms reference "last year", questionnaire partial
  PACKRT  packright_quote.pdf     ex-works + settlement discount + index clause in footnotes
  CORRUG  corrugated_ratecard.jpg angled photo; 3-ply per-100, 5/7-ply PER CASE with the
                                  case-pack note smudged (pack-factor gate) + one blurred cell
  GLOBAL  globalpack_email.txt    7-ply in USD, "freight included" in body vs csv attachment
          globalpack_rates.csv    saying ex-works/freight extra → source conflict

Run:  python -m backend.make_vendor_files      (deterministic — same files every run)
"""
from __future__ import annotations

import csv
import io
from pathlib import Path

from .seed_data import SKUS, _weight_kg

OUT = Path(__file__).resolve().parent.parent / "data" / "vendor_files"
KG_RATE = {3: 37.5, 5: 41.0, 7: 45.0}
CASE_PACK = 102
FX = 83.60


def _jit(key: str, amp: float = 0.02) -> float:
    h = 2166136261
    for ch in key:
        h = ((h ^ ord(ch)) * 16777619) & 0xFFFFFFFF
    return 1 + ((h % 10000) / 10000 * 2 - 1) * amp


def base_pc(code, L, W, H, ply, gsm) -> float:
    return _weight_kg(L, W, H, gsm) * KG_RATE[ply]


ROWS = [(c, L, W, H, ply, gsm, prn, qty) for c, L, W, H, ply, gsm, prn, qty, _ in SKUS]
DESC = {c: f"RSC {L}x{W}x{H} {ply}P" + ("-PRINT" if prn else "") for c, L, W, H, ply, gsm, prn, q in ROWS}
PRINTED = [c for c, *_ , prn, q in [(r[0], r[1], r[2], r[3], r[4], r[5], r[6], r[7]) for r in ROWS] if _ is None]  # unused guard


def printed_codes():
    return [r[0] for r in ROWS if r[6]]


# ---------------- ALPHA (xlsx, clean control) ----------------

def make_alpha():
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "Quote"
    ws.append(["Item Code", "Description", "Monthly Qty", "Rate (INR per piece)", "Remarks"])
    for c, L, W, H, ply, gsm, prn, qty in ROWS:
        rate = round(base_pc(c, L, W, H, ply, gsm) * 1.00 * _jit("ALPHA" + c), 2)
        ws.append([c, DESC[c], qty, rate, "FOR Bhiwandi, GST extra"])
    q = wb.create_sheet("Compliance")
    q.append(["Question", "Answer"])
    for row in [
        ["ISO 9001:2015", "Yes - cert AQP-3312, valid to Jun 2028 (attached)"],
        ["Burst strength certs (IS 2771)", "Yes - all quoted grades, attached"],
        ["FSC chain of custody", "Yes - FSC-C140221"],
        ["Lead time (days)", 8],
        ["45-day payment accepted", "Yes"],
        ["Capacity headroom (%)", 35],
    ]:
        q.append(row)
    wb.save(OUT / "alpha_quote.xlsx")


# ---------------- BOXCO (docx, 27/30 + references) ----------------

def make_boxco():
    import docx
    d = docx.Document()
    d.add_heading("BoxCo Containers Pvt Ltd — Commercial Offer", level=1)
    d.add_paragraph(
        "Ref: RFQ-FY27-CORR. We thank you for the enquiry and are pleased to quote for the "
        "plain corrugated range on a delivered-Bhiwandi basis, GST extra. Our rates per piece "
        "are tabulated below.")
    t = d.add_table(rows=1, cols=3)
    hdr = t.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "Item", "Description", "Rate (Rs / piece)"
    for c, L, W, H, ply, gsm, prn, qty in ROWS:
        if prn:
            continue
        f = 1.035 if ply == 3 else (0.995 if ply == 5 else 1.01)
        rate = round(base_pc(c, L, W, H, ply, gsm) * f * _jit("BOXCO" + c), 2)
        r = t.add_row().cells
        r[0].text, r[1].text, r[2].text = c, DESC[c], f"{rate:.2f}"
    d.add_paragraph(
        "Please note we regret our inability to undertake the printed cartons "
        "(PKG-006, PKG-013 and PKG-023) as we do not operate printing lines.")
    d.add_paragraph(
        "Commercial terms: payment terms shall remain the same as last year's arrangement. "
        "Prices firm for 12 months. We hold ISO 9001:2015 (certificate QMS-2214, valid to "
        "March 2028) and our standard lead time is 9 days from confirmed order.")
    d.save(OUT / "boxco_quote.docx")


# ---------------- PACKRIGHT (pdf, footnotes) ----------------

def make_packright():
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.pdfgen import canvas
    path = str(OUT / "packright_quote.pdf")
    cv = canvas.Canvas(path, pagesize=A4)
    wpt, hpt = A4

    def header(title):
        cv.setFont("Helvetica-Bold", 14)
        cv.drawString(20 * mm, hpt - 20 * mm, "PACKRIGHT INDUSTRIES")
        cv.setFont("Helvetica", 8)
        cv.drawString(20 * mm, hpt - 25 * mm, "Plot 44, Chakan MIDC, Pune | GSTIN 27AAPCP1180F1Z1 | Ref PR/Q/2617")
        cv.setFont("Helvetica-Bold", 10)
        cv.drawString(20 * mm, hpt - 33 * mm, title)
        cv.line(20 * mm, hpt - 35 * mm, wpt - 20 * mm, hpt - 35 * mm)

    header("Quotation — RFQ-FY27-CORR (rates in Rs per piece)")
    cv.setFont("Helvetica", 8.5)
    y = hpt - 42 * mm
    for i, (c, L, W, H, ply, gsm, prn, qty) in enumerate(ROWS):
        rate = round(base_pc(c, L, W, H, ply, gsm) * 0.962 * _jit("PACKRT" + c), 2)
        cv.drawString(22 * mm, y, c)
        cv.drawString(45 * mm, y, DESC[c])
        cv.drawRightString(120 * mm, y, f"{rate:.2f} *")
        y -= 5 * mm
        if y < 40 * mm and i < len(ROWS) - 1:
            cv.showPage()
            header("Quotation (contd.)")
            cv.setFont("Helvetica", 8.5)
            y = hpt - 42 * mm
    # questionnaire block
    y -= 6 * mm
    cv.setFont("Helvetica-Bold", 9)
    cv.drawString(20 * mm, y, "Compliance declarations")
    cv.setFont("Helvetica", 8.5)
    for line in [
        "ISO 9001:2015 - Yes, cert PRI-9917 valid to Nov 2027 (copy enclosed).",
        "Burst strength test certificates (IS 2771) - enclosed for all quoted grades.",
        "FSC chain of custody - Yes, FSC-C121077.",
        "Lead time - 11 days. 45-day payment terms - accepted. Capacity headroom - approx 28%.",
    ]:
        y -= 4.5 * mm
        cv.drawString(22 * mm, y, line)
    # the footnotes that matter
    y -= 9 * mm
    cv.setFont("Helvetica-Oblique", 7.5)
    for fn in [
        "*  Prices are ex-works Chakan. Freight payable at actuals.",
        "** A settlement discount of 2.5% applies to payments made within 10 days of invoice.",
        "†  Prices subject to revision should the kraft paper index vary beyond ±5% during the contract period.",
    ]:
        y -= 4 * mm
        cv.drawString(20 * mm, y, fn)
    cv.save()


# ---------------- CORRUGATED (angled photo rate card) ----------------

def make_corrugated():
    from PIL import Image, ImageDraw, ImageFilter, ImageFont

    W_, H_ = 1500, 2000
    card = Image.new("RGB", (W_, H_), (233, 221, 195))
    dr = ImageDraw.Draw(card)

    def font(sz, bold=False):
        for name in (["/System/Library/Fonts/Supplemental/Courier New Bold.ttf"] if bold else
                     ["/System/Library/Fonts/Supplemental/Courier New.ttf"]):
            try:
                return ImageFont.truetype(name, sz)
            except OSError:
                continue
        return ImageFont.load_default()

    f_h, f_b, f_s = font(44, True), font(30), font(26)
    CW = dr.textlength("M", font=f_s)  # mono char width at row size
    dr.text((W_ // 2, 60), "CORRUGATED INDUSTRIES", font=f_h, anchor="mm", fill=(60, 45, 25))
    dr.text((W_ // 2, 110), "WORKS RATE CARD  FY 2026-27  (FOR BHIWANDI, GST EXTRA)",
            font=f_s, anchor="mm", fill=(60, 45, 25))
    dr.line((80, 140, W_ - 80, 140), fill=(120, 100, 60), width=3)

    blur_boxes = []
    X0 = 100

    def row_line(c, L, W, H, prn, rate_txt, y, blur_rate=False):
        left = f"{c}  {L}x{W}x{H}" + ("  PRINTED" if prn else "")
        pad = 46 - len(left) - len(rate_txt)
        line = left + " " + "." * max(pad, 2) + " " + rate_txt
        dr.text((X0, y), line, font=f_s, fill=(60, 45, 25))
        if blur_rate:
            i = line.rindex(rate_txt)
            blur_boxes.append((int(X0 + i * CW) - 6, y - 6, int(X0 + (i + len(rate_txt)) * CW) + 8, y + 34))

    y = 190
    dr.text((X0, y), "3-PLY CARTONS  -  RATE PER 100 PCS", font=f_b, fill=(60, 45, 25)); y += 48
    for c, L, W, H, ply, gsm, prn, qty in ROWS:
        if ply != 3:
            continue
        per100 = round(base_pc(c, L, W, H, ply, gsm) * 0.99 * _jit("CORRUG" + c) * 100 / 5) * 5
        row_line(c, L, W, H, prn, f"Rs {per100:,}", y)
        y += 40
    y += 30
    dr.text((X0, y), "5-PLY / 7-PLY  -  RATE PER CASE", font=f_b, fill=(60, 45, 25)); y += 48
    for c, L, W, H, ply, gsm, prn, qty in ROWS:
        if ply == 3:
            continue
        pc = base_pc(c, L, W, H, ply, gsm) * 0.985 * _jit("CORRUG" + c)
        per_case = round(pc * CASE_PACK / 10) * 10
        row_line(c, L, W, H, prn, f"Rs {per_case:,}", y, blur_rate=(c == "PKG-022"))
        y += 40
    y += 36
    note = f"NOTE: STD CASE PACK = {CASE_PACK} PCS. PLAIN+PRINTED BOTH OK. PAYMENT 30 DAYS."
    dr.text((X0, y), note, font=f_s, fill=(60, 45, 25))
    i = note.index(str(CASE_PACK))
    blur_boxes.append((int(X0 + (i - 1) * CW), y - 6, int(X0 + (i + 4) * CW) + 6, y + 32))
    dr.text((X0, y + 44), "CONTACT: 98220-XXXXX  (RAJESH)", font=f_s, fill=(60, 45, 25))

    for bx in blur_boxes:
        region = card.crop(bx).filter(ImageFilter.GaussianBlur(6))
        card.paste(region, bx)

    photo = card.rotate(2.8, expand=True, fillcolor=(72, 66, 55), resample=Image.BICUBIC)
    stage = Image.new("RGB", (photo.width + 120, photo.height + 120), (58, 53, 44))
    stage.paste(photo, (60, 60))
    stage = stage.filter(ImageFilter.GaussianBlur(0.6))
    stage.save(OUT / "corrugated_ratecard.jpg", quality=88)


# ---------------- GLOBALPACK (email + csv attachment, conflict) ----------------

def make_globalpack():
    usd = {}
    inr = {}
    for c, L, W, H, ply, gsm, prn, qty in ROWS:
        if ply == 7:
            usd[c] = round(base_pc(c, L, W, H, ply, gsm) * 0.952 * _jit("GLOBAL" + c) / FX, 2)
        else:
            inr[c] = round(base_pc(c, L, W, H, ply, gsm) * 1.02 * _jit("GLOBAL" + c), 2)
    body = f"""From: Dinesh Mehta <dinesh@globalpack.co>
To: rfq-fy27@parity.suraksha.in
Date: Mon, 31 Aug 2026 21:12:44 +0530
Subject: RE: RFQ FY27 - corrugated rates
Attachment: globalpack_rates.csv

Madam,

Sharing our offer. For the heavy-duty 7-ply export cartons:
PKG-027 at USD {usd['PKG-027']:.2f}, PKG-028 at USD {usd['PKG-028']:.2f},
PKG-029 at USD {usd['PKG-029']:.2f}, PKG-030 at USD {usd['PKG-030']:.2f} per piece.

All other sizes as per the attached rate list. Freight is included up to your
Bhiwandi DC. 45 days payment is fine.

Regards,
Dinesh Mehta
GlobalPack Corrugators, Daman
"""
    (OUT / "globalpack_email.txt").write_text(body)
    buf = io.StringIO()
    w = csv.writer(buf)
    w.writerow(["item_code", "description", "rate_inr_per_piece"])
    for c in sorted(inr):
        w.writerow([c, DESC[c], f"{inr[c]:.2f}"])
    w.writerow([])
    w.writerow(["Note: all rates ex-works Daman; freight at actuals.", "", ""])
    (OUT / "globalpack_rates.csv").write_text(buf.getvalue())


VENDORS = {
    "ALPHA": {"name": "Alpha Packaging", "files": ["alpha_quote.xlsx"],
              "contact": "Suresh Iyer", "email": "rates@alphapackaging.in"},
    "BOXCO": {"name": "BoxCo Containers", "files": ["boxco_quote.docx"],
              "contact": "Meera Shah", "email": "sales@boxco.co.in"},
    "PACKRT": {"name": "PackRight Industries", "files": ["packright_quote.pdf"],
               "contact": "Anil Deshpande", "email": "quotes@packright.in"},
    "CORRUG": {"name": "Corrugated Industries", "files": ["corrugated_ratecard.jpg"],
               "contact": "Rajesh", "email": "corrugatedind@gmail.com"},
    "GLOBAL": {"name": "GlobalPack Corrugators", "files": ["globalpack_email.txt", "globalpack_rates.csv"],
               "contact": "Dinesh Mehta", "email": "dinesh@globalpack.co"},
}


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    make_alpha(); make_boxco(); make_packright(); make_corrugated(); make_globalpack()
    import json
    (OUT.parent / "vendors.json").write_text(json.dumps(VENDORS, indent=1))
    for f in sorted(OUT.iterdir()):
        print(f"  {f.name}  {f.stat().st_size:,} bytes")


if __name__ == "__main__":
    main()
